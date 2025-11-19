#!/usr/bin/env python3
"""
Practical example of using content merging for enhanced bullet generation.

This demonstrates how the new content merging methods can be used to
generate better bullets by providing table context.
"""

import sys
import os
from docx import Document

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from slide_generator_pkg.document_parser import DocumentParser


def create_realistic_document():
    """Create a realistic business document with tables and context"""
    doc = Document()

    # Title
    doc.add_heading('Q4 2024 Sales Performance Report', level=1)

    # Executive summary
    doc.add_paragraph(
        'This report summarizes sales performance for Q4 2024. '
        'All regions showed strong growth compared to the previous quarter.'
    )

    # Section 1: Regional Performance
    doc.add_heading('Regional Sales Performance', level=2)

    # Context before table
    doc.add_paragraph(
        'The following table shows sales figures by region. North America led '
        'with the highest revenue, while APAC showed the strongest growth rate.'
    )

    # Sales table
    table1 = doc.add_table(rows=5, cols=4)
    table1.style = 'Light Grid Accent 1'

    # Headers
    headers = ['Region', 'Revenue ($M)', 'Growth (%)', 'Market Share (%)']
    for idx, header in enumerate(headers):
        table1.rows[0].cells[idx].text = header

    # Data
    data = [
        ['North America', '125.4', '12.5', '42'],
        ['Europe', '89.2', '8.3', '30'],
        ['APAC', '65.8', '23.1', '22'],
        ['Latin America', '18.6', '15.7', '6']
    ]

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, value in enumerate(row_data):
            table1.rows[row_idx].cells[col_idx].text = value

    # Context after table
    doc.add_paragraph(
        'Key takeaway: APAC growth of 23.1% represents a strategic opportunity '
        'for expansion in 2025. North America maintains market leadership despite '
        'moderate growth.'
    )

    # Section 2: Product Performance
    doc.add_heading('Top Products by Category', level=2)

    # Standalone table (no context paragraphs)
    table2 = doc.add_table(rows=4, cols=3)
    table2.style = 'Light Grid Accent 1'

    table2.rows[0].cells[0].text = 'Product'
    table2.rows[0].cells[1].text = 'Category'
    table2.rows[0].cells[2].text = 'Units Sold'

    products = [
        ['CloudSync Pro', 'Software', '45,000'],
        ['DataVault Enterprise', 'Hardware', '12,500'],
        ['AIAssist Premium', 'SaaS', '78,000']
    ]

    for idx, product in enumerate(products, start=1):
        for col_idx, value in enumerate(product):
            table2.rows[idx].cells[col_idx].text = value

    # Save document
    file_path = '/tmp/realistic_sales_report.docx'
    doc.save(file_path)
    print(f"Created realistic document: {file_path}\n")
    return file_path


def demonstrate_content_merging():
    """Demonstrate the content merging workflow"""
    print("=" * 80)
    print("Content Merging Demonstration - Practical Use Case")
    print("=" * 80)
    print()

    # Create realistic document
    doc_path = create_realistic_document()

    # Initialize parser
    parser = DocumentParser()

    # Load document
    doc = Document(doc_path)

    print("STEP 1: Extract Content Blocks")
    print("-" * 80)

    # Extract content blocks
    content_blocks = parser._extract_content_blocks_from_docx(doc)
    print(f"Extracted {len(content_blocks)} content blocks\n")

    for idx, block in enumerate(content_blocks):
        print(f"{idx + 1}. {block['type'].upper()}", end='')
        if block['type'] == 'heading':
            print(f" (Level {block['level']}): {block['text'][:60]}...")
        elif block['type'] == 'paragraph':
            print(f": {block['text'][:60]}...")
        elif block['type'] == 'table':
            print(f": {len(block['data'])} rows x {len(block['data'][0]) if block['data'] else 0} columns")

    print("\n" + "=" * 80)
    print("STEP 2: Merge Tables with Context")
    print("-" * 80)

    # Merge tables with context
    merged_blocks = parser._merge_table_and_text_context(content_blocks)
    print(f"Merged into {len(merged_blocks)} blocks\n")

    for idx, block in enumerate(merged_blocks):
        print(f"\nBlock {idx + 1}: {block['type'].upper()}")

        if block['type'] == 'table_with_context':
            print("  Context Analysis:")
            if block['intro']:
                print(f"    ✓ Intro: {block['intro'][:70]}...")
            else:
                print("    ✗ No intro context")

            print(f"    • Table: {len(block['table']['data'])} rows")

            if block['explanation']:
                print(f"    ✓ Explanation: {block['explanation'][:70]}...")
            else:
                print("    ✗ No explanation context")

        elif block['type'] == 'heading':
            print(f"  Level {block['level']}: {block['text']}")

        elif block['type'] == 'paragraph':
            print(f"  Text: {block['text'][:70]}...")

    print("\n" + "=" * 80)
    print("STEP 3: Demonstrate Bullet Generation with Context")
    print("-" * 80)

    for idx, block in enumerate(merged_blocks):
        if block['type'] == 'table_with_context':
            print(f"\nTable {idx + 1}:")

            # Build context for bullet generation
            context_parts = []
            if block['intro']:
                context_parts.append(f"INTRO: {block['intro']}")
            if block['explanation']:
                context_parts.append(f"SUMMARY: {block['explanation']}")

            if context_parts:
                print("\n  Combined Context for Bullet Generation:")
                print(f"  {' | '.join(context_parts)}")
                print("\n  → This context would enhance LLM bullet generation")
                print("  → Bullets would be more accurate and relevant")
            else:
                print("\n  No context available")
                print("  → Would use table data only for bullet generation")

    print("\n" + "=" * 80)
    print("COMPARISON: With vs Without Context")
    print("-" * 80)

    print("\nWITHOUT Context Merging:")
    print("  Bullets would be generated from table data alone:")
    print("  • Region: North America, Revenue: $125.4M, Growth: 12.5%")
    print("  • Region: Europe, Revenue: $89.2M, Growth: 8.3%")
    print("  • Region: APAC, Revenue: $65.8M, Growth: 23.1%")

    print("\nWITH Context Merging:")
    print("  Bullets incorporate surrounding context:")
    print("  • North America leads with $125.4M revenue (42% market share)")
    print("  • APAC shows strongest growth at 23.1%, representing strategic")
    print("    opportunity for 2025 expansion")
    print("  • Europe maintains steady performance with $89.2M revenue")

    print("\n" + "=" * 80)
    print("✓ Demonstration Complete")
    print("=" * 80)
    print("\nBenefits of Content Merging:")
    print("  1. More contextual and meaningful bullets")
    print("  2. Incorporates author's interpretations and insights")
    print("  3. Better alignment with document narrative")
    print("  4. Improved audience comprehension")
    print("  5. More professional presentation quality")
    print()


if __name__ == '__main__':
    demonstrate_content_merging()
