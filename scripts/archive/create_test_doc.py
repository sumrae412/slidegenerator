#!/usr/bin/env python3
"""
Create a test document for learner outcome testing
"""
from docx import Document

# Create a new document
doc = Document()

# Add title
doc.add_heading('Streamlit Development Course', level=1)

# Add Module 1
doc.add_heading('Module 1: Introduction to Streamlit', level=2)

p1 = doc.add_paragraph(
    "Streamlit is a powerful Python library that allows developers to create interactive web applications with minimal code. "
    "Unlike traditional web frameworks, Streamlit focuses on simplicity and rapid prototyping. "
    "Developers can build data visualization apps, machine learning dashboards, and interactive tools without needing extensive web development knowledge."
)

# Add subheading
doc.add_heading('Key Features and Benefits', level=3)

p2 = doc.add_paragraph(
    "The framework provides automatic reactivity, meaning the app updates when users interact with widgets or when underlying data changes. "
    "It supports various input widgets like sliders, text inputs, and file uploaders. "
    "The library integrates seamlessly with popular data science libraries including pandas, matplotlib, and plotly."
)

# Add another subheading  
doc.add_heading('Building Your First App', level=4)

p3 = doc.add_paragraph(
    "Creating a Streamlit app involves writing Python scripts with special Streamlit commands. "
    "The st.write() function displays content, while st.sidebar creates interactive elements. "
    "Developers can deploy apps locally or to cloud platforms for sharing with others."
)

# Add Module 2
doc.add_heading('Module 2: Advanced Streamlit Techniques', level=2)

p4 = doc.add_paragraph(
    "Advanced Streamlit development includes state management, custom components, and performance optimization. "
    "Session state allows apps to maintain data across user interactions. "
    "Custom CSS styling can improve app appearance, while caching decorators enhance performance for data-heavy applications."
)

# Save the document
doc.save('test_learner_outcomes.docx')
print("Created test_learner_outcomes.docx")