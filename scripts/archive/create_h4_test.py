#!/usr/bin/env python3
"""
Create a test DOCX document with H4 headings to test the slide generation
"""

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

# Create a new document
doc = Document()

# Add title
title = doc.add_heading('Test Presentation with H4 Headings', level=1)

# Add H2 section
section_h2 = doc.add_heading('Section: Learning Concepts', level=2)

# Add H3 subsection  
subsection_h3 = doc.add_heading('Subsection: Core Principles', level=3)

# Add H4 headings (individual slide titles)
h4_1 = doc.add_heading('What is Machine Learning', level=4)
doc.add_paragraph('Machine learning is a subset of artificial intelligence that enables computers to learn and improve from experience.')

h4_2 = doc.add_heading('How to Implement Neural Networks', level=4) 
doc.add_paragraph('Neural networks can be implemented using various frameworks and require careful attention to architecture design.')

h4_3 = doc.add_heading('Key Benefits of Deep Learning', level=4)
doc.add_paragraph('Deep learning offers significant advantages in pattern recognition, automation, and predictive analytics.')

# Add another H2 section
section_h2_2 = doc.add_heading('Section: Practical Applications', level=2)

# Add more H4 headings
h4_4 = doc.add_heading('Best Practices for Data Preparation', level=4)
doc.add_paragraph('Proper data preparation is crucial for successful machine learning model training and validation.')

h4_5 = doc.add_heading('Common Challenges in AI Development', level=4)
doc.add_paragraph('AI development faces challenges including data quality, computational resources, and ethical considerations.')

# Save the document
doc.save('test_h4_document.docx')
print("Created test_h4_document.docx with H4 headings for testing")