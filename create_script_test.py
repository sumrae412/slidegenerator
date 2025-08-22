from docx import Document

# Create a new Document
doc = Document()

# Add title
title = doc.add_heading('Video Script Document', 0)

# Add section 1
doc.add_heading('Module 1: Introduction', level=1)

# Add a script table
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

# Add header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Video'
hdr_cells[1].text = 'Script Content'
hdr_cells[2].text = 'Notes'

# Add script data rows
script_data = [
    ('V1', 'Welcome to our course on rapid prototyping. In this video, we will explore the fundamentals of building applications quickly using modern tools and frameworks. You will learn how to move from idea to working prototype in just a few hours.', 'Opening hook'),
    ('V2', 'Before we dive into the technical details, let us understand why prototyping matters. Traditional development takes weeks or months, but with the right approach, you can validate ideas in days. This is especially important in today\'s fast-moving tech landscape.', 'Context setting'),
    ('V3', 'GenAI has revolutionized how we build software. Instead of writing every line of code from scratch, you can describe what you want and get working code instantly. This means faster iteration, quicker testing, and more time for creativity.', 'Technology intro'),
    ('V4', 'Let me show you a practical example. Suppose you want to build a sentiment analysis tool. Traditionally, this would require setting up libraries, writing parsing logic, and creating a user interface. With GenAI, you simply ask for what you need.', 'Practical demo'),
    ('V5', 'The key to successful prototyping is starting small and iterating quickly. Do not try to build everything at once. Focus on one core feature, test it with real users, and then expand based on feedback. This approach reduces risk and increases your chances of success.', 'Best practices')
]

for video, script, notes in script_data:
    row_cells = table.add_row().cells
    row_cells[0].text = video
    row_cells[1].text = script
    row_cells[2].text = notes

# Add section 2
doc.add_heading('Module 2: Advanced Topics', level=1)

# Add another script table
table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'

# Add header row
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'Video'
hdr_cells2[1].text = 'Script Content'
hdr_cells2[2].text = 'Duration'

# Add more script data
advanced_data = [
    ('V6', 'Now that you understand the basics, let us explore more advanced prototyping techniques. We will cover testing strategies, user feedback collection, and how to scale your prototypes into production applications.', '5 min'),
    ('V7', 'Testing is crucial for any prototype. You need to validate not just that your code works, but that it solves the right problem. This means getting it in front of real users as quickly as possible and gathering meaningful feedback.', '4 min'),
    ('V8', 'As you collect feedback, you will need to iterate on your design. This is where the agile mindset becomes essential. Be prepared to pivot, modify, or even completely rebuild based on what you learn from users.', '6 min')
]

for video, script, duration in advanced_data:
    row_cells = table2.add_row().cells
    row_cells[0].text = video
    row_cells[1].text = script
    row_cells[2].text = duration

# Save the document
doc.save('script_test_document.docx')
print("Created script_test_document.docx with script content")