#!/usr/bin/env python3
"""
Create test document with the exact Streamlit example to verify bullet point style
"""

from docx import Document

# Create a new document
doc = Document()

# Add title
doc.add_heading('Streamlit Development Test', level=1)

# Create a table with the exact example content
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'

# Set headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Time'
hdr_cells[1].text = 'Script Content'

# Add the exact example paragraph
row1_cells = table.rows[1].cells
row1_cells[0].text = '0:00'
row1_cells[1].text = '''Streamlit allows developers to rapidly create interactive web apps with just a few lines of Python code. Its intuitive API simplifies the process of adding widgets, visualizing data, and integrating machine learning models. This makes it especially useful for prototyping AI tools and dashboards, where speed and iteration are critical.'''

# Save the document
doc.save('streamlit_test_document.docx')
print("Created streamlit_test_document.docx with the exact example paragraph")