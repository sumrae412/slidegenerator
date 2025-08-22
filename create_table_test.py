#!/usr/bin/env python3
"""
Create a test DOCX document with table content to test bullet point generation
"""

from docx import Document

# Create a new document
doc = Document()

# Add title
doc.add_heading('Machine Learning Presentation Test', level=1)

# Create a table with script content
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'

# Set headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Time'
hdr_cells[1].text = 'Script Content'

# Add substantial content rows
row1_cells = table.rows[1].cells
row1_cells[0].text = '0:00'
row1_cells[1].text = '''Machine learning is revolutionizing how we approach data analysis and problem-solving in the modern world. This technology enables computers to learn patterns from data without being explicitly programmed for each specific task. The main benefit is that machine learning systems can improve their performance over time as they are exposed to more data. For example, recommendation systems on streaming platforms use machine learning to suggest content based on viewing history. The process involves three key stages: data preprocessing, model training, and performance evaluation. Companies like Netflix and Amazon have seen significant improvements in customer satisfaction because machine learning helps them understand user preferences better.'''

row2_cells = table.rows[2].cells
row2_cells[0].text = '2:30'
row2_cells[1].text = '''Deep learning represents a subset of machine learning that uses artificial neural networks with multiple layers to model complex patterns. The main advantage of deep learning is its ability to automatically discover intricate features in data without manual feature engineering. This approach has led to breakthrough results in computer vision, natural language processing, and speech recognition. For instance, image classification systems can now identify objects in photographs with accuracy that rivals human performance. The training process requires large amounts of labeled data and significant computational power, typically using specialized hardware like GPUs. While deep learning offers impressive capabilities, it also presents challenges such as the need for extensive datasets, long training times, and difficulty in interpreting model decisions.'''

# Save the document
doc.save('table_test_document.docx')
print("Created table_test_document.docx with substantial table content for testing")