#!/usr/bin/env python3
"""
Create a test document without tables for paragraph-based slide generation
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_paragraph_test_doc():
    """Create a test document with only paragraphs (no tables)"""
    doc = Document()
    
    # Add title (H1)
    title = doc.add_heading('Machine Learning Fundamentals', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add section (H2)
    doc.add_heading('Introduction to ML', level=2)
    
    # Add paragraphs that will become slides
    doc.add_paragraph(
        "Machine learning is a subset of artificial intelligence that enables computers to learn from data without being explicitly programmed. "
        "It involves algorithms that improve their performance through experience, identifying patterns and making decisions with minimal human intervention. "
        "The field has revolutionized industries from healthcare to finance."
    )
    
    doc.add_paragraph(
        "There are three main types of machine learning: supervised learning, unsupervised learning, and reinforcement learning. "
        "Supervised learning uses labeled data to train models for prediction tasks. "
        "Unsupervised learning finds hidden patterns in unlabeled data. "
        "Reinforcement learning involves agents learning through trial and error."
    )
    
    # Add another section (H2)
    doc.add_heading('Neural Networks', level=2)
    
    doc.add_paragraph(
        "Neural networks are computing systems inspired by biological neural networks in animal brains. "
        "They consist of interconnected nodes called neurons organized in layers. "
        "Information flows from input layer through hidden layers to output layer. "
        "Deep learning uses neural networks with many hidden layers."
    )
    
    doc.add_paragraph(
        "Training neural networks involves forward propagation and backpropagation. "
        "Forward propagation passes input through the network to generate predictions. "
        "Backpropagation calculates gradients to update weights and minimize error. "
        "This iterative process continues until the model converges."
    )
    
    # Add subsection (H3)
    doc.add_heading('Applications', level=3)
    
    doc.add_paragraph(
        "Computer vision uses neural networks for image classification, object detection, and facial recognition. "
        "Convolutional neural networks excel at processing visual information. "
        "Applications include autonomous vehicles, medical imaging, and security systems."
    )
    
    doc.add_paragraph(
        "Natural language processing enables machines to understand and generate human language. "
        "Transformer models like GPT and BERT have achieved breakthrough performance. "
        "Applications include chatbots, translation, sentiment analysis, and text generation."
    )
    
    # Save the document
    doc.save('paragraph_test_document.docx')
    print("✅ Created paragraph_test_document.docx")
    print("📄 Document contains:")
    print("   - 1 H1 heading")
    print("   - 2 H2 headings")
    print("   - 1 H3 heading")
    print("   - 6 content paragraphs")
    print("   - 0 tables")
    print("\n🎯 Test by selecting 'No table - Use paragraphs for slides' option")

if __name__ == "__main__":
    create_paragraph_test_doc()