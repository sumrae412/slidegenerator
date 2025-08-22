#!/usr/bin/env python3
"""
Create a test document with actual content for visual prompts
"""
from docx import Document

# Create a new document
doc = Document()

# Add title
doc.add_heading('GenAI App Development Course', level=1)

# Add content paragraphs that will generate content slides
doc.add_paragraph(
    "Welcome to GenAI app development! You'll learn to build interactive applications using Streamlit and OpenAI. "
    "This course focuses on rapid prototyping and iterative development. "
    "By the end, you'll have deployed a working GenAI application to the cloud."
)

doc.add_paragraph(
    "Prototyping with GenAI requires a different mindset than traditional development. "
    "You'll learn to prompt AI models effectively, handle unpredictable outputs, and create user interfaces quickly. "
    "The goal is to test ideas fast and iterate based on real user feedback."
)

doc.add_paragraph(
    "Streamlit makes it easy to build web applications with Python. "
    "You can create interactive widgets, display data visualizations, and deploy to the cloud with just a few lines of code. "
    "This framework is perfect for GenAI prototypes because it handles the web development complexity for you."
)

# Save the document
doc.save('content_test_document.docx')
print("Created content_test_document.docx with actual content for visual prompts")
