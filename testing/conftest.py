"""
Pytest configuration and shared fixtures for slide generator tests
"""

import pytest
import os
import sys
import tempfile
import shutil
from unittest.mock import Mock, patch
from io import BytesIO

# Add the parent directory to the path so we can import our modules
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from file_to_slides_enhanced import (
    app, DocumentProcessor, PresentationGenerator, 
    ProcessingOptions, SlideContent, TEMPLATES
)
from docx import Document
import PyPDF2

@pytest.fixture
def flask_app():
    """Flask app fixture for testing"""
    app.config['TESTING'] = True
    app.config['WTF_CSRF_ENABLED'] = False
    return app

@pytest.fixture
def client(flask_app):
    """Flask test client"""
    return flask_app.test_client()

@pytest.fixture
def temp_dir():
    """Temporary directory for test files"""
    temp_dir = tempfile.mkdtemp()
    yield temp_dir
    shutil.rmtree(temp_dir)

@pytest.fixture
def processing_options():
    """Default processing options for tests"""
    return ProcessingOptions(
        template='professional',
        bullet_density=3,
        include_visual_prompts=True,
        generate_speaker_notes=False,
        add_slide_numbers=True,
        slides_per_section=5,
        column_index=0,
        api_key=None
    )

@pytest.fixture
def processing_options_with_api():
    """Processing options with API key for tests"""
    return ProcessingOptions(
        template='creative',
        bullet_density=4,
        include_visual_prompts=True,
        generate_speaker_notes=True,
        add_slide_numbers=True,
        slides_per_section=3,
        column_index=0,
        api_key='sk-test-api-key-12345'
    )

@pytest.fixture
def sample_text_content():
    """Sample text content for testing"""
    return {
        'paragraphs': [
            "This is the first paragraph with important information about the topic.",
            "The second paragraph contains additional details and context.",
            "Finally, the third paragraph provides conclusions and next steps.",
            "Advanced features include AI-powered bullet generation and visual prompts.",
            "The system supports multiple file formats and template options."
        ],
        'headings': [
            {'level': 1, 'text': 'Main Presentation Title'},
            {'level': 2, 'text': 'Introduction Section'},
            {'level': 3, 'text': 'Key Features'},
            {'level': 4, 'text': 'Implementation Details'}
        ],
        'tables': [
            [
                ['Column 1', 'Column 2', 'Column 3'],
                ['Data 1A', 'Data 2A', 'Data 3A'],
                ['Data 1B', 'Data 2B', 'Data 3B'],
                ['Data 1C', 'Data 2C', 'Data 3C']
            ]
        ]
    }

@pytest.fixture
def sample_slide_content():
    """Sample slide content for testing"""
    return [
        SlideContent(
            title="Introduction",
            bullets=[
                "Welcome to our presentation.",
                "This covers key topics and insights.",
                "We'll explore advanced features."
            ],
            visual_prompt="Visual concept: introduction, presentation - modern diagram",
            speaker_notes="Introduce the topic and set expectations",
            slide_type='content',
            template_style=TEMPLATES['professional']
        ),
        SlideContent(
            title="Key Features",
            bullets=[
                "AI-powered content generation.",
                "Multiple template options available.",
                "Real-time progress tracking.",
                "Secure file processing."
            ],
            visual_prompt="Visual concept: features, technology - professional illustration",
            speaker_notes="Highlight the main features and benefits",
            slide_type='content',
            template_style=TEMPLATES['professional']
        )
    ]

@pytest.fixture
def create_test_docx_file(temp_dir):
    """Create a test DOCX file"""
    def _create_docx(filename="test.docx", content=None):
        doc = Document()
        
        # Add title
        title = doc.add_heading('Test Presentation', 1)
        
        # Add section heading
        doc.add_heading('Introduction Section', 2)
        
        # Add subsection
        doc.add_heading('Key Points', 3)
        
        # Add individual slide heading
        doc.add_heading('First Topic', 4)
        
        # Add paragraphs
        if content:
            for para in content.get('paragraphs', []):
                doc.add_paragraph(para)
        else:
            doc.add_paragraph("This is a test paragraph with important information.")
            doc.add_paragraph("Another paragraph with additional details.")
            doc.add_paragraph("Final paragraph with conclusions.")
        
        # Add a table
        table = doc.add_table(rows=3, cols=3)
        table_data = [
            ['Script Column', 'Notes', 'Timing'],
            ['First slide content here', 'Important note', '2 min'],
            ['Second slide content here', 'Another note', '3 min']
        ]
        
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = cell_data
        
        filepath = os.path.join(temp_dir, filename)
        doc.save(filepath)
        return filepath
    
    return _create_docx

@pytest.fixture
def create_test_txt_file(temp_dir):
    """Create a test TXT file"""
    def _create_txt(filename="test.txt", content=None):
        filepath = os.path.join(temp_dir, filename)
        
        default_content = """# Main Presentation Title

## Introduction Section

### Key Features

#### Implementation Details

This is the first paragraph with important information about the topic.

The second paragraph contains additional details and context.

Finally, the third paragraph provides conclusions and next steps.

#### Advanced Topics

Advanced features include AI-powered bullet generation and visual prompts.

The system supports multiple file formats and template options.

## Conclusion

Thank you for your attention. Questions are welcome.
"""
        
        with open(filepath, 'w', encoding='utf-8') as f:
            if content:
                f.write(content)
            else:
                f.write(default_content)
        
        return filepath
    
    return _create_txt

@pytest.fixture
def create_test_pdf_file(temp_dir):
    """Create a test PDF file (simplified for testing)"""
    def _create_pdf(filename="test.pdf", content=None):
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        filepath = os.path.join(temp_dir, filename)
        
        c = canvas.Canvas(filepath, pagesize=letter)
        width, height = letter
        
        # Add title
        c.setFont("Helvetica-Bold", 24)
        c.drawString(72, height - 72, "Test Presentation")
        
        # Add content
        c.setFont("Helvetica", 12)
        y = height - 120
        
        lines = [
            "Introduction Section",
            "",
            "This is a test document for PDF processing.",
            "It contains multiple paragraphs and sections.",
            "The content will be converted to presentation slides.",
            "",
            "Key Features:",
            "- AI-powered processing",
            "- Multiple format support", 
            "- Template customization",
            "",
            "Thank you for testing our system."
        ]
        
        for line in lines:
            if line:
                c.drawString(72, y, line)
            y -= 20
            if y < 72:  # Start new page if needed
                c.showPage()
                c.setFont("Helvetica", 12)
                y = height - 72
        
        c.save()
        return filepath
    
    return _create_pdf

@pytest.fixture
def mock_openai_response():
    """Mock OpenAI API response"""
    return Mock(
        choices=[
            Mock(
                message=Mock(
                    content="""• First bullet point about the topic
• Second bullet point with more details
• Third bullet point summarizing key insights
• Fourth bullet point providing conclusions"""
                )
            )
        ]
    )

@pytest.fixture
def session_id():
    """Test session ID"""
    return "test-session-12345"

# Performance test fixtures
@pytest.fixture
def large_text_content():
    """Large text content for performance testing"""
    paragraphs = []
    for i in range(100):
        paragraphs.append(f"This is paragraph {i+1} with detailed information about topic {i+1}. " * 10)
    
    headings = []
    for i in range(20):
        headings.append({'level': 2, 'text': f'Section {i+1}'})
        headings.append({'level': 3, 'text': f'Subsection {i+1}'})
        headings.append({'level': 4, 'text': f'Topic {i+1}'})
    
    return {
        'paragraphs': paragraphs,
        'headings': headings,
        'tables': []
    }

# Error testing fixtures
@pytest.fixture
def invalid_content():
    """Invalid content for error testing"""
    return {
        'paragraphs': [],
        'headings': [],
        'tables': []
    }

@pytest.fixture
def malformed_content():
    """Malformed content for error testing"""
    return {
        'paragraphs': [None, "", "   ", "Valid paragraph"],
        'headings': [
            {'level': 0, 'text': ''},
            {'level': 10, 'text': 'Invalid level'},
            {'level': 2, 'text': 'Valid heading'}
        ],
        'tables': [[]]
    }