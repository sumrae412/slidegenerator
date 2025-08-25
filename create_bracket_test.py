#!/usr/bin/env python3
"""
Create a test document with bracketed stage directions to test filtering
"""
from docx import Document

def create_bracket_test_doc():
    """Create a test document with various bracketed stage directions"""
    doc = Document()
    
    # Add title
    doc.add_heading('Stage Direction Test Document', level=1)
    
    # Add paragraphs with various bracket types
    doc.add_paragraph(
        "[SCREENCAST] Welcome to this presentation about artificial intelligence. [CLICK] "
        "We'll explore how AI is transforming industries across the globe. [PAUSE:3] "
        "This technology has applications in healthcare, finance, and transportation."
    )
    
    doc.add_paragraph(
        "[VIDEO-DEMO] Let's start with machine learning basics. [SLIDE_TRANSITION] "
        "Machine learning algorithms learn from data to make predictions. [CLICK] "
        "There are three main types: supervised, unsupervised, and reinforcement learning. [PAUSE]"
    )
    
    doc.add_paragraph(
        "[ANNOTATION] Deep learning is a subset of machine learning. [HIGHLIGHT] "
        "It uses neural networks with multiple layers to process information. [ZOOM_IN] "
        "These networks can recognize patterns in images, text, and audio."
    )
    
    # Add paragraph with mixed case brackets (should be preserved)
    doc.add_paragraph(
        "Some text with [lowercase brackets] should remain. [UPPER-CASE] should be removed. "
        "This sentence has [Mixed Case] brackets and [ALL_CAPS] brackets for testing."
    )
    
    # Add table to test table processing
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # Header row
    table.cell(0, 0).text = "Topic"
    table.cell(0, 1).text = "Script Content"
    
    # Content rows with stage directions
    table.cell(1, 0).text = "Introduction"
    table.cell(1, 1).text = "[SCREENCAST] Neural networks are inspired by the human brain. [CLICK] They consist of interconnected nodes called neurons."
    
    table.cell(2, 0).text = "Applications"
    table.cell(2, 1).text = "[DEMO] Computer vision uses neural networks for image recognition. [PAUSE:2] Applications include medical imaging and autonomous vehicles."
    
    # Save document
    doc.save('bracket_test_document.docx')
    print("✅ Created bracket_test_document.docx")
    print("\n📄 Document contains:")
    print("   - Title with H1")
    print("   - 4 paragraphs with stage directions")
    print("   - 1 table with stage directions")
    print("   - Mixed case brackets for testing")
    print("\n🎯 Stage directions to be filtered:")
    print("   - [SCREENCAST], [CLICK], [PAUSE:3]")
    print("   - [VIDEO-DEMO], [SLIDE_TRANSITION]")
    print("   - [ANNOTATION], [HIGHLIGHT], [ZOOM_IN]")
    print("   - [UPPER-CASE], [ALL_CAPS]")
    print("\n💡 Should preserve: [lowercase brackets], [Mixed Case]")

if __name__ == "__main__":
    create_bracket_test_doc()