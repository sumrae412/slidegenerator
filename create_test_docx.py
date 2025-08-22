from docx import Document
from docx.shared import Inches

# Create a new Document
doc = Document()

# Add title
title = doc.add_heading('Test Document for Slides', 0)

# Add section 1
doc.add_heading('Section 1: Employee Data', level=1)
doc.add_heading('Overview', level=2)
doc.add_paragraph('This section contains employee information.')

# Add a table
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'

# Add header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Name'
hdr_cells[1].text = 'Age'
hdr_cells[2].text = 'Department'
hdr_cells[3].text = 'Salary'

# Add data rows
data = [
    ('John Smith', '25', 'Engineering', '75000'),
    ('Jane Doe', '30', 'Marketing', '65000'),
    ('Bob Johnson', '28', 'Sales', '55000'),
    ('Alice Brown', '35', 'HR', '60000')
]

for name, age, dept, salary in data:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = age
    row_cells[2].text = dept
    row_cells[3].text = salary

# Add section 2
doc.add_heading('Section 2: Product Data', level=1)
doc.add_paragraph('Product inventory information.')

# Add another table
table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'

# Add header row
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'Product'
hdr_cells2[1].text = 'Price'
hdr_cells2[2].text = 'Stock'

# Add product data
products = [
    ('Laptop', '1200', '50'),
    ('Mouse', '25', '200'),
    ('Keyboard', '75', '150'),
    ('Monitor', '300', '80')
]

for product, price, stock in products:
    row_cells = table2.add_row().cells
    row_cells[0].text = product
    row_cells[1].text = price
    row_cells[2].text = stock

# Save the document
doc.save('test_document.docx')
print("Created test_document.docx with tables")